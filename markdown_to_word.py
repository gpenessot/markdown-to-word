import os
import re
import json
import argparse
import tempfile
import requests
from urllib.parse import urlparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.section import WD_HEADER_FOOTER
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from bs4 import BeautifulSoup
from PIL import Image
import io

# Configuration par défaut
DEFAULT_CONFIG = {
    "styles": {
        "h1": {
            "font_name": "Arial",
            "font_size": 18,
            "bold": True,
            "color": {"r": 0, "g": 51, "b": 102},
            "space_before": 24,
            "space_after": 12,
            "keep_with_next": True
        },
        "h2": {
            "font_name": "Arial",
            "font_size": 16,
            "bold": True,
            "color": {"r": 0, "g": 102, "b": 153},
            "space_before": 18,
            "space_after": 10,
            "keep_with_next": True
        },
        "h3": {
            "font_name": "Arial",
            "font_size": 14,
            "bold": True,
            "color": {"r": 0, "g": 153, "b": 204},
            "space_before": 14,
            "space_after": 8,
            "keep_with_next": True
        },
        "h4": {
            "font_name": "Arial",
            "font_size": 12,
            "bold": True,
            "color": {"r": 0, "g": 153, "b": 204},
            "space_before": 12,
            "space_after": 6,
            "keep_with_next": True
        },
        "normal": {
            "font_name": "Times New Roman",
            "font_size": 11,
            "color": {"r": 0, "g": 0, "b": 0},
            "space_after": 8,
            "line_spacing": 1.15
        },
        "code": {
            "font_name": "Consolas",
            "font_size": 10,
            "color": {"r": 128, "g": 0, "b": 128},
            "space_after": 8
        },
        "code_block": {
            "font_name": "Consolas",
            "font_size": 10,
            "color": {"r": 0, "g": 0, "b": 0},
            "background": {"r": 245, "g": 245, "b": 245},
            "space_before": 8,
            "space_after": 8
        },
        "list_item": {
            "font_name": "Times New Roman",
            "font_size": 11,
            "color": {"r": 0, "g": 0, "b": 0},
            "space_after": 0,
            "left_indent": 0.5
        },
        "table": {
            "font_name": "Times New Roman",
            "font_size": 10,
            "header_bg_color": {"r": 240, "g": 240, "b": 240},
            "border_color": {"r": 0, "g": 0, "b": 0},
            "space_before": 12,
            "space_after": 12
        },
        "caption": {
            "font_name": "Times New Roman",
            "font_size": 10,
            "italic": True,
            "color": {"r": 80, "g": 80, "b": 80},
            "space_before": 6,
            "space_after": 12,
            "alignment": "center"
        },
        "toc_heading": {
            "font_name": "Arial",
            "font_size": 16,
            "bold": True,
            "color": {"r": 0, "g": 0, "b": 0},
            "space_before": 24,
            "space_after": 12
        },
        "toc_item": {
            "font_name": "Times New Roman",
            "font_size": 11,
            "color": {"r": 0, "g": 0, "b": 0},
            "space_after": 3
        }
    },
    "document": {
        "add_file_headers": True,
        "page_break_between_files": True,
        "margins": {
            "top": 2.5,   # en cm
            "bottom": 2.5,
            "left": 2.5,
            "right": 2.5
        },
        "page_size": {
            "width": 21.0,  # A4 en cm
            "height": 29.7
        },
        "orientation": "portrait",
        "generate_toc": True,
        "toc_title": "Table des matières",
        "header": {
            "enabled": True,
            "content": "{chapter}",
            "first_page_different": True
        },
        "footer": {
            "enabled": True,
            "content": "Page {page}",
            "first_page_different": True
        },
        "image_max_width": 15  # en cm
    }
}

def load_config(config_file=None):
    """
    Charge la configuration depuis un fichier JSON
    
    Args:
        config_file (str): Chemin vers le fichier de configuration JSON
        
    Returns:
        dict: Configuration chargée ou configuration par défaut
    """
    config = DEFAULT_CONFIG.copy()
    
    if config_file and os.path.exists(config_file):
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                user_config = json.load(f)
                
            # Fusion récursive des configurations
            merge_configs(config, user_config)
                
        except Exception as e:
            print(f"Erreur lors du chargement de la configuration: {e}")
            print("Utilisation de la configuration par défaut.")
    
    return config

def merge_configs(default_config, user_config):
    """
    Fusionne récursivement la configuration utilisateur avec la configuration par défaut
    
    Args:
        default_config (dict): Configuration par défaut
        user_config (dict): Configuration utilisateur
    """
    for key, value in user_config.items():
        if key in default_config and isinstance(default_config[key], dict) and isinstance(value, dict):
            merge_configs(default_config[key], value)
        else:
            default_config[key] = value

def extract_image_caption(element):
    """
    Extrait la légende d'une image depuis le texte alt ou title
    
    Args:
        element (Tag): Élément BeautifulSoup image
        
    Returns:
        str: Légende de l'image ou None
    """
    caption = element.get('alt') or element.get('title')
    return caption

def download_image(url, temp_dir=None):
    """
    Télécharge une image depuis une URL et la sauvegarde dans un dossier temporaire
    
    Args:
        url (str): URL de l'image à télécharger
        temp_dir (str): Dossier temporaire (si None, utilise le dossier temp du système)
        
    Returns:
        str: Chemin vers l'image téléchargée, None en cas d'erreur
    """
    if temp_dir is None:
        temp_dir = tempfile.gettempdir()
    
    try:
        # Créer un nom de fichier à partir de l'URL
        parsed_url = urlparse(url)
        filename = os.path.basename(parsed_url.path)
        
        # S'assurer que le nom de fichier est valide
        if not filename or '.' not in filename:
            filename = f"image_{hash(url)}.jpg"
        
        # Chemin complet du fichier
        filepath = os.path.join(temp_dir, filename)
        
        # Télécharger l'image
        response = requests.get(url, stream=True, timeout=10)
        response.raise_for_status()  # Vérifier si la requête a réussi
        
        # Sauvegarder l'image
        with open(filepath, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        return filepath
    
    except Exception as e:
        print(f"Erreur lors du téléchargement de l'image {url}: {e}")
        return None

def add_image_to_doc(doc, img_path, base_path, config, caption=None):
    """
    Ajoute une image au document Word
    
    Args:
        doc (Document): Document Word
        img_path (str): Chemin de l'image (peut être relatif ou URL)
        base_path (str): Chemin de base pour les chemins relatifs
        config (dict): Configuration
        caption (str): Légende de l'image
    """
    # Vérifier si c'est une URL
    is_url = img_path.startswith(('http://', 'https://'))
    
    if is_url:
        # Tenter de télécharger l'image
        print(f"Téléchargement de l'image distante: {img_path}")
        downloaded_img_path = download_image(img_path)
        
        if downloaded_img_path and os.path.exists(downloaded_img_path):
            img_path = downloaded_img_path
        else:
            # Si le téléchargement a échoué
            p = doc.add_paragraph(f"[Image externe non téléchargée: {img_path}]")
            apply_para_style(p, config["styles"]["normal"])
            return
    else:
        # Gestion des chemins relatifs pour les fichiers locaux
        if not os.path.isabs(img_path):
            img_path = os.path.join(base_path, img_path)
        
        if not os.path.exists(img_path):
            print(f"Avertissement: Impossible de trouver l'image {img_path}")
            p = doc.add_paragraph(f"[Image non trouvée: {os.path.basename(img_path)}]")
            apply_para_style(p, config["styles"]["normal"])
            return
    
    try:
        # Ajout de l'image
        max_width = Cm(config["document"]["image_max_width"])
        
        # Redimensionnement intelligent de l'image pour qu'elle s'adapte à la page
        img = Image.open(img_path)
        width, height = img.size
        
        # Convertir max_width de cm à pixels (approximation : 1 cm = 37.8 pixels)
        max_width_px = max_width.cm * 37.8
        
        # Si l'image est plus large que la largeur maximale, redimensionner
        if width > max_width_px:
            ratio = max_width_px / width
            new_width = max_width_px
            # Conserver le ratio d'aspect
            height = height * ratio
        else:
            new_width = width
        
        # Ajouter l'image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(new_width / 37.8))
        
        # Ajouter la légende si présente
        if caption:
            cap_p = doc.add_paragraph(f"Figure: {caption}")
            apply_para_style(cap_p, config["styles"]["caption"])
    except Exception as e:
        print(f"Erreur lors de l'ajout de l'image {img_path}: {e}")
        p = doc.add_paragraph(f"[Erreur lors du chargement de l'image: {os.path.basename(img_path)}]")
        apply_para_style(p, config["styles"]["normal"])

def add_table_to_doc(doc, table_html, config):
    """
    Ajoute un tableau au document Word à partir du HTML
    
    Args:
        doc (Document): Document Word
        table_html (Tag): Tableau HTML (BeautifulSoup)
        config (dict): Configuration
    """
    # Extraire les lignes et colonnes
    rows = table_html.find_all('tr')
    if not rows:
        return
    
    # Déterminer le nombre de colonnes en fonction de la première ligne
    header_row = rows[0]
    header_cells = header_row.find_all(['th', 'td'])
    num_cols = len(header_cells)
    
    # Créer le tableau Word
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'  # Style de base avec bordures
    
    # Parcourir les lignes et colonnes
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        
        for j, cell in enumerate(cells):
            if j < num_cols:  # Éviter les index hors limites
                # Obtenir le texte de la cellule
                cell_text = cell.get_text(strip=True)
                
                # Appliquer le texte à la cellule
                word_cell = table.cell(i, j)
                word_cell.text = cell_text
                
                # Appliquer le style
                for paragraph in word_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = config["styles"]["table"]["font_name"]
                        run.font.size = Pt(config["styles"]["table"]["font_size"])
                
                # Style pour les en-têtes (première ligne)
                if i == 0 or cell.name == 'th':
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    
                    # Couleur de fond pour les en-têtes
                    header_bg = config["styles"]["table"]["header_bg_color"]
                    shade_cell(word_cell, header_bg["r"], header_bg["g"], header_bg["b"])
    
    # Espace après le tableau
    doc.add_paragraph()

def shade_cell(cell, r, g, b):
    """
    Ajoute une couleur de fond à une cellule de tableau
    
    Args:
        cell: Cellule de tableau Word
        r, g, b: Valeurs RGB pour la couleur
    """
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), f"{r:02x}{g:02x}{b:02x}")
    cell._tc.get_or_add_tcPr().append(shading)

def add_list_to_doc(doc, list_html, config, list_type=None, level=0):
    """
    Ajoute une liste (ordonnée ou non) au document Word
    
    Args:
        doc (Document): Document Word
        list_html (Tag): Élément de liste HTML (ol/ul)
        config (dict): Configuration
        list_type (str): Type de liste ('ol' ou 'ul')
        level (int): Niveau d'indentation
    """
    if not list_type:
        list_type = list_html.name  # 'ol' ou 'ul'
    
    items = list_html.find_all('li', recursive=False)
    
    for item in items:
        # Texte de l'élément (sans les sous-listes)
        item_text = ''
        for content in item.contents:
            if content.name not in ['ol', 'ul']:
                if hasattr(content, 'get_text'):
                    item_text += content.get_text()
                else:
                    item_text += str(content)
        
        item_text = item_text.strip()
        
        # Créer un paragraphe pour l'élément de liste
        p = doc.add_paragraph()
        p.style = 'List Bullet' if list_type == 'ul' else 'List Number'
        
        # Ajuster l'indentation en fonction du niveau
        p.paragraph_format.left_indent = Inches((level + 1) * float(config["styles"]["list_item"]["left_indent"]))
        
        # Ajouter le texte de l'élément
        run = p.add_run(item_text)
        apply_style(run, config["styles"]["list_item"])
        
        # Traiter les sous-listes récursivement
        for sub_list in item.find_all(['ol', 'ul'], recursive=False):
            add_list_to_doc(doc, sub_list, config, sub_list.name, level + 1)

def add_header_footer(section, header_text, footer_text, config, chapter_title):
    """
    Ajoute des en-têtes et pieds de page à la section
    
    Args:
        section: Section Word
        header_text (str): Texte de l'en-tête
        footer_text (str): Texte du pied de page
        config (dict): Configuration
        chapter_title (str): Titre du chapitre
    """
    # En-tête
    if config["document"]["header"]["enabled"]:
        header = section.header
        
        # S'assurer qu'il y a au moins un paragraphe
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
            p.clear()  # Vider le contenu existant
        
        # Ajouter le contenu
        p.text = header_text.replace("{chapter}", chapter_title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pied de page
    if config["document"]["footer"]["enabled"]:
        footer = section.footer
        
        # S'assurer qu'il y a au moins un paragraphe
        if not footer.paragraphs:
            p = footer.add_paragraph()
        else:
            p = footer.paragraphs[0]
            p.clear()  # Vider le contenu existant
        
        # Ajouter le texte statique
        p.text = footer_text.replace("{page}", "")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajouter le numéro de page en utilisant un champ
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def add_section_with_settings(doc, config, chapter_title):
    """
    Ajoute une nouvelle section avec les paramètres de page configurés
    
    Args:
        doc (Document): Document Word
        config (dict): Configuration
        chapter_title (str): Titre du chapitre actuel
        
    Returns:
        section: Nouvelle section ajoutée
    """
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Orientation
    if config["document"]["orientation"].lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    
    # Taille de page
    section.page_width = Cm(config["document"]["page_size"]["width"])
    section.page_height = Cm(config["document"]["page_size"]["height"])
    
    # Marges
    section.top_margin = Cm(config["document"]["margins"]["top"])
    section.bottom_margin = Cm(config["document"]["margins"]["bottom"])
    section.left_margin = Cm(config["document"]["margins"]["left"])
    section.right_margin = Cm(config["document"]["margins"]["right"])
    
    # En-têtes et pieds de page
    if config["document"]["header"]["enabled"] or config["document"]["footer"]["enabled"]:
        # Page différente pour la première page si configurée
        section.different_first_page_header_footer = (
            config["document"]["header"]["first_page_different"] or 
            config["document"]["footer"]["first_page_different"]
        )
        
        header_text = config["document"]["header"]["content"]
        footer_text = config["document"]["footer"]["content"]
        
        add_header_footer(section, header_text, footer_text, config, chapter_title)
    
    return section

def generate_toc(doc, config):
    """
    Génère une table des matières
    
    Args:
        doc (Document): Document Word
        config (dict): Configuration
    """
    if not config["document"]["generate_toc"]:
        return
    
    print("Génération de la table des matières...")
    
    # Ajouter un titre pour la table des matières
    heading = doc.add_heading(config["document"]["toc_title"], level=1)
    apply_heading_style(heading, config["styles"]["toc_heading"])
    
    # Ajouter un paragraphe pour la TOC
    par = doc.add_paragraph()
    
    # Méthode plus stable pour ajouter un champ TOC
    run = par.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z'
    run._r.append(instr)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    # Texte par défaut qui sera remplacé lors de la mise à jour
    run = par.add_run("Table des matières (Cliquez-droit et sélectionnez 'Mettre à jour les champs' pour générer)")
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # Ajouter une page après la TOC
    doc.add_page_break()
    
    print("Table des matières ajoutée. Ouvrez le document et mettez à jour la table pour l'afficher correctement.")


def apply_style(run, style):
    """
    Applique un style à un élément Run
    
    Args:
        run: L'élément Run (texte) à styliser
        style (dict): Dictionnaire contenant les propriétés de style
    """
    if "font_name" in style:
        run.font.name = style["font_name"]
        
    if "font_size" in style:
        run.font.size = Pt(style["font_size"])
        
    if "bold" in style:
        run.bold = style["bold"]
        
    if "italic" in style:
        run.italic = style["italic"]
        
    if "color" in style:
        r = style["color"].get("r", 0)
        g = style["color"].get("g", 0)
        b = style["color"].get("b", 0)
        run.font.color.rgb = RGBColor(r, g, b)

def apply_para_style(paragraph, style):
    """
    Applique un style à un paragraphe
    
    Args:
        paragraph: Le paragraphe à styliser
        style (dict): Dictionnaire contenant les propriétés de style
    """
    # Appliquer le style à chaque Run dans le paragraphe
    for run in paragraph.runs:
        apply_style(run, style)
    
    # Appliquer les propriétés spécifiques aux paragraphes
    if "space_before" in style:
        paragraph.paragraph_format.space_before = Pt(style["space_before"])
        
    if "space_after" in style:
        paragraph.paragraph_format.space_after = Pt(style["space_after"])
        
    if "line_spacing" in style:
        paragraph.paragraph_format.line_spacing = style["line_spacing"]
        
    if "keep_with_next" in style:
        paragraph.paragraph_format.keep_with_next = style["keep_with_next"]
        
    if "alignment" in style:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if style["alignment"] in alignment_map:
            paragraph.paragraph_format.alignment = alignment_map[style["alignment"]]

def apply_heading_style(heading, style):
    """
    Applique un style à un élément Heading
    
    Args:
        heading: L'élément Heading à styliser
        style (dict): Dictionnaire contenant les propriétés de style
    """
    apply_para_style(heading, style)
    
    for run in heading.runs:
        apply_style(run, style)

def extract_chapter_title(md_content):
    """
    Extrait le titre du chapitre (premier titre h1) d'un contenu Markdown
    
    Args:
        md_content (str): Contenu Markdown
        
    Returns:
        str: Titre du chapitre ou nom de fichier par défaut
    """
    # Chercher le premier titre h1
    h1_match = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
    if h1_match:
        return h1_match.group(1).strip()
    return "Chapitre"

def convert_markdown_to_docx(input_files, output_file, config_file=None):
    """
    Convertit un ou plusieurs fichiers Markdown en un document Word (.docx)
    
    Args:
        input_files (list): Liste des chemins vers les fichiers Markdown
        output_file (str): Chemin du fichier Word de sortie
        config_file (str): Chemin du fichier de configuration
    """
    # Charger la configuration
    config = load_config(config_file)
    
    # Créer un nouveau document Word
    doc = Document()
    
    # Appliquer les paramètres de page initiaux
    first_section = doc.sections[0]
    first_section.page_width = Cm(config["document"]["page_size"]["width"])
    first_section.page_height = Cm(config["document"]["page_size"]["height"])
    first_section.top_margin = Cm(config["document"]["margins"]["top"])
    first_section.bottom_margin = Cm(config["document"]["margins"]["bottom"])
    first_section.left_margin = Cm(config["document"]["margins"]["left"])
    first_section.right_margin = Cm(config["document"]["margins"]["right"])
    
    # Générer la table des matières en premier si demandé
    if config["document"]["generate_toc"]:
        # Configuration de la section pour la table des matières
        toc_section = first_section
        toc_section.different_first_page_header_footer = (
            config["document"]["header"]["first_page_different"] or 
            config["document"]["footer"]["first_page_different"]
        )
        
        # Appliquer l'en-tête et le pied de page à la section de la table des matières
        if config["document"]["header"]["enabled"] or config["document"]["footer"]["enabled"]:
            header_text = config["document"]["header"]["content"]
            footer_text = config["document"]["footer"]["content"]
            add_header_footer(toc_section, header_text, footer_text, config, config["document"]["toc_title"])
        
        # Ajouter un titre pour la table des matières
        heading = doc.add_heading(config["document"]["toc_title"], level=1)
        apply_heading_style(heading, config["styles"]["toc_heading"])
        
        # Ajouter un paragraphe pour la TOC
        par = doc.add_paragraph()
        
        # Méthode plus stable pour ajouter un champ TOC
        run = par.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-3" \\h \\z'
        run._r.append(instr)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar)
        
        # Texte par défaut qui sera remplacé lors de la mise à jour
        run = par.add_run("Table des matières (Cliquez-droit et sélectionnez 'Mettre à jour les champs' pour générer)")
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        
        # Ajouter une page après la TOC
        doc.add_page_break()
        
        print("Table des matières ajoutée. Ouvrez le document et mettez à jour la table pour l'afficher correctement.")
    
    # Traiter chaque fichier Markdown
    for index, input_file in enumerate(input_files):
        # Lire le contenu du fichier Markdown
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Extraire le titre du chapitre
        chapter_title = extract_chapter_title(md_content)
        file_name = os.path.basename(input_file)
        
        print(f"Traitement du fichier: {file_name} (chapitre: {chapter_title})")
        
        # Ajouter une nouvelle section pour chaque fichier (sauf pour le tout premier si TOC n'est pas générée)
        if index > 0 or config["document"]["generate_toc"]:
            # Ajouter un saut de page entre les chapitres si nécessaire
            if config["document"]["page_break_between_files"]:
                doc.add_page_break()
            
            # Créer une nouvelle section
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            
            # Configurer les paramètres de page pour cette section
            section.page_width = Cm(config["document"]["page_size"]["width"])
            section.page_height = Cm(config["document"]["page_size"]["height"])
            section.top_margin = Cm(config["document"]["margins"]["top"])
            section.bottom_margin = Cm(config["document"]["margins"]["bottom"])
            section.left_margin = Cm(config["document"]["margins"]["left"])
            section.right_margin = Cm(config["document"]["margins"]["right"])
            
            # Configurer l'option pour différencier la première page
            section.different_first_page_header_footer = (
                config["document"]["header"]["first_page_different"] or 
                config["document"]["footer"]["first_page_different"]
            )
        else:
            # Utiliser la première section déjà existante
            section = doc.sections[0]
        
        # Appliquer immédiatement les en-têtes et pieds de page pour cette section
        if config["document"]["header"]["enabled"] or config["document"]["footer"]["enabled"]:
            header_text = config["document"]["header"]["content"]
            footer_text = config["document"]["footer"]["content"]
            add_header_footer(section, header_text, footer_text, config, chapter_title)
        
        # Obtenir le chemin de base pour les images relatives
        base_path = os.path.dirname(os.path.abspath(input_file))
        
        # Ajouter le nom du fichier comme titre (optionnel)
        if config["document"]["add_file_headers"]:
            file_header = doc.add_heading(f"Fichier: {file_name}", level=1)
            apply_heading_style(file_header, config["styles"]["h1"])
        
        # Traitement direct du Markdown
        lines = md_content.split('\n')
        i = 0
        
        in_code_block = False
        code_language = None
        code_content = []
        
        while i < len(lines):
            line = lines[i]
            
            # Détection des blocs de code avec triple backticks
            if line.strip().startswith('```'):
                if not in_code_block:
                    # Début du bloc de code
                    in_code_block = True
                    # Récupérer le langage (ex: ```python)
                    code_language = line.strip()[3:].strip()
                    code_content = []
                else:
                    # Fin du bloc de code
                    in_code_block = False
                    
                    # Ajouter le bloc de code au document
                    if code_language:
                        lang_p = doc.add_paragraph()
                        lang_run = lang_p.add_run(f"Code ({code_language}):")
                        lang_run.bold = True
                    
                    # Ajouter le contenu du code
                    code_text = '\n'.join(code_content)
                    code_p = doc.add_paragraph()
                    code_run = code_p.add_run(code_text)
                    
                    # Appliquer le style au bloc de code
                    apply_style(code_run, config["styles"]["code_block"])
                    
                    # Ajouter un espace après le bloc de code
                    doc.add_paragraph()
                i += 1
                continue
            
            if in_code_block:
                # Ajouter cette ligne au contenu du code
                code_content.append(line)
                i += 1
                continue
            
            # Traitement des titres (h1 à h6)
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                title_text = header_match.group(2).strip()
                
                # Ajouter le titre au document
                heading = doc.add_heading(title_text, level=level)
                
                # Appliquer un style spécifique
                style_key = f'h{level}'
                if style_key in config["styles"]:
                    apply_heading_style(heading, config["styles"][style_key])
                
                i += 1
                continue
            
            # Traitement des éléments de liste
            list_match = re.match(r'^(\s*)([\*\-\+]|\d+\.)\s+(.+)$', line)
            if list_match:
                indentation = len(list_match.group(1))
                list_marker = list_match.group(2)
                content = list_match.group(3)
                
                # Déterminer le type de liste
                if list_marker in ['*', '-', '+']:
                    current_list_type = 'ul'
                else:
                    current_list_type = 'ol'
                
                # Calculer le niveau d'indentation
                indent_level = indentation // 2
                
                # Créer un paragraphe pour l'élément de liste
                p = doc.add_paragraph()
                p.style = 'List Bullet' if current_list_type == 'ul' else 'List Number'
                p.paragraph_format.space_after = Pt(0)  # Réduire l'espace après chaque élément
                
                # Ajuster l'indentation en fonction du niveau
                p.paragraph_format.left_indent = Inches((indent_level + 1) * float(config["styles"]["list_item"]["left_indent"]))
                
                # Traiter le contenu avec formatage inline
                remaining_text = content
                formatted_runs = []
                
                # Analyser tout le texte pour identifier les portions formatées
                while remaining_text:
                    # Texte en gras: **texte** ou __texte__
                    bold_match = re.search(r'(\*\*|__)(.*?)(\1)', remaining_text)
                    # Texte en italique: *texte* ou _texte_
                    italic_match = re.search(r'([*_])((?!\1).*?)(\1)', remaining_text)
                    # Code inline: `texte`
                    code_match = re.search(r'`(.*?)`', remaining_text)
                    # Lien: [texte](url)
                    link_match = re.search(r'\[(.*?)\]\((.*?)\)', remaining_text)
                    
                    # Trouver le premier match
                    matches = [(bold_match, 'bold'), (italic_match, 'italic'), 
                            (code_match, 'code'), (link_match, 'link')]
                    valid_matches = [(m, t) for m, t in matches if m]
                    
                    if valid_matches:
                        # Trier par position de début
                        valid_matches.sort(key=lambda x: x[0].start())
                        match, match_type = valid_matches[0]
                        
                        # Ajouter le texte avant le match
                        if match.start() > 0:
                            formatted_runs.append(('normal', remaining_text[:match.start()]))
                        
                        # Ajouter le match avec son style
                        if match_type == 'bold':
                            formatted_runs.append(('bold', match.group(2)))
                        elif match_type == 'italic':
                            formatted_runs.append(('italic', match.group(2)))
                        elif match_type == 'code':
                            formatted_runs.append(('code', match.group(1)))
                        elif match_type == 'link':
                            formatted_runs.append(('link', match.group(1)))
                        
                        # Continuer avec le reste du texte
                        remaining_text = remaining_text[match.end():]
                    else:
                        # Aucun format spécial trouvé, ajouter tout le texte restant
                        formatted_runs.append(('normal', remaining_text))
                        break
                
                # Ajouter tous les runs au paragraphe avec leur format
                for format_type, text in formatted_runs:
                    run = p.add_run(text)
                    if format_type == 'bold':
                        run.bold = True
                    elif format_type == 'italic':
                        run.italic = True
                    elif format_type == 'link':
                        run.underline = True
                    
                    # Appliquer le style approprié
                    if format_type == 'code':
                        apply_style(run, config["styles"]["code"])
                    else:
                        apply_style(run, config["styles"]["list_item"])
                
                i += 1
                continue
            
            # Traitement des tableaux (syntaxe markdown)
            if line.strip().startswith('|') and line.strip().endswith('|'):
                # Collecte des lignes de tableau
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                
                # S'assurer qu'il y a au moins un séparateur et une ligne de données
                if len(table_lines) >= 3:
                    # Extraction des cellules
                    header_row = [cell.strip() for cell in table_lines[0].strip('|').split('|')]
                    rows_data = []
                    
                    # Ignorer la ligne de séparation (---|---|...)
                    for table_line in table_lines[2:]:
                        rows_data.append([cell.strip() for cell in table_line.strip('|').split('|')])
                    
                    # Créer le tableau Word
                    num_cols = len(header_row)
                    table = doc.add_table(rows=len(rows_data) + 1, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    # Remplir l'en-tête
                    for j, cell_text in enumerate(header_row):
                        cell = table.cell(0, j)
                        cell.text = cell_text
                        
                        # Style d'en-tête
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.name = config["styles"]["table"]["font_name"]
                        
                        # Couleur de fond
                        header_bg = config["styles"]["table"]["header_bg_color"]
                        shade_cell(cell, header_bg["r"], header_bg["g"], header_bg["b"])
                    
                    # Remplir les données
                    for row_idx, row_data in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:  # Éviter les index hors limites
                                cell = table.cell(row_idx + 1, col_idx)
                                cell.text = cell_text
                                
                                # Style de cellule
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = config["styles"]["table"]["font_name"]
                                        run.font.size = Pt(config["styles"]["table"]["font_size"])
                    
                    # Espace après le tableau
                    doc.add_paragraph()
                continue
            
            # Détection d'images markdown ![alt](url)
            img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)
                add_image_to_doc(doc, img_path, base_path, config, alt_text)
                i += 1
                continue
            
            # Traitement du texte normal
            # Si ligne vide ou seulement des espaces
            if not line.strip():
                doc.add_paragraph()
                i += 1
                continue
            
            # Texte normal
            p = doc.add_paragraph()
            
            # Traitement du texte avec formatage inline amélioré
            remaining_text = line
            formatted_runs = []
            
            # Analyser tout le texte pour identifier les portions formatées
            while remaining_text:
                # Texte en gras: **texte** ou __texte__
                bold_match = re.search(r'(\*\*|__)(.*?)(\1)', remaining_text)
                # Texte en italique: *texte* ou _texte_
                italic_match = re.search(r'([*_])((?!\1).*?)(\1)', remaining_text)
                # Code inline: `texte`
                code_match = re.search(r'`(.*?)`', remaining_text)
                # Lien: [texte](url)
                link_match = re.search(r'\[(.*?)\]\((.*?)\)', remaining_text)
                
                # Trouver le premier match
                matches = [(bold_match, 'bold'), (italic_match, 'italic'), 
                        (code_match, 'code'), (link_match, 'link')]
                valid_matches = [(m, t) for m, t in matches if m]
                
                if valid_matches:
                    # Trier par position de début
                    valid_matches.sort(key=lambda x: x[0].start())
                    match, match_type = valid_matches[0]
                    
                    # Ajouter le texte avant le match
                    if match.start() > 0:
                        formatted_runs.append(('normal', remaining_text[:match.start()]))
                    
                    # Ajouter le match avec son style
                    if match_type == 'bold':
                        formatted_runs.append(('bold', match.group(2)))
                    elif match_type == 'italic':
                        formatted_runs.append(('italic', match.group(2)))
                    elif match_type == 'code':
                        formatted_runs.append(('code', match.group(1)))
                    elif match_type == 'link':
                        formatted_runs.append(('link', match.group(1)))
                    
                    # Continuer avec le reste du texte
                    remaining_text = remaining_text[match.end():]
                else:
                    # Aucun format spécial trouvé, ajouter tout le texte restant
                    formatted_runs.append(('normal', remaining_text))
                    break
            
            # Ajouter tous les runs au paragraphe avec leur format
            for format_type, text in formatted_runs:
                run = p.add_run(text)
                if format_type == 'bold':
                    run.bold = True
                elif format_type == 'italic':
                    run.italic = True
                elif format_type == 'link':
                    run.underline = True
                
                # Appliquer le style approprié
                if format_type == 'code':
                    apply_style(run, config["styles"]["code"])
                else:
                    apply_style(run, config["styles"]["normal"])
            
            # Appliquer le style de paragraphe
            apply_para_style(p, config["styles"]["normal"])
            
            i += 1
    
    # Enregistrer le document Word
    doc.save(output_file)
    print(f"Conversion terminée. Document Word enregistré sous: {output_file}")


def create_default_config(output_file="config.json"):
    """
    Crée un fichier de configuration par défaut
    
    Args:
        output_file (str): Chemin du fichier de configuration à créer
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(DEFAULT_CONFIG, f, indent=4)
    
    print(f"Fichier de configuration créé: {output_file}")

def main():
    parser = argparse.ArgumentParser(description='Convertir des fichiers Markdown en document Word avec support avancé.')
    parser.add_argument('input', nargs='+', help='Fichier(s) Markdown à convertir')
    parser.add_argument('-o', '--output', default='output.docx', help='Fichier Word de sortie')
    parser.add_argument('-c', '--config', help='Fichier de configuration (JSON)')
    parser.add_argument('--create-config', action='store_true', help='Créer un fichier de configuration par défaut')
    
    args = parser.parse_args()
    
    if args.create_config:
        create_default_config()
        return
    
    convert_markdown_to_docx(args.input, args.output, args.config)

if __name__ == "__main__":
    main()